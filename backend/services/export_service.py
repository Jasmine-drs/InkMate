"""
导出服务
支持 TXT、EPUB、DOCX 格式导出
"""
import io
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import AsyncIterator

from ebooklib import epub
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func

from models.chapter import Chapter
from models.project import Project
from models.unit import Unit


class ExportService:
    """导出服务"""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def get_project_chapters(self, project_id: str) -> list[Chapter]:
        """获取项目的所有章节（按章节号排序）"""
        result = await self.db.execute(
            select(Chapter)
            .where(Chapter.project_id == project_id)
            .order_by(Chapter.chapter_number.asc())
        )
        return list(result.scalars().all())

    async def get_project_info(self, project_id: str) -> Project | None:
        """获取项目信息"""
        result = await self.db.execute(
            select(Project).where(Project.id == project_id)
        )
        return result.scalar_one_or_none()

    async def get_units_map(self, project_id: str) -> dict[str, Unit]:
        """获取项目的单元映射表"""
        result = await self.db.execute(
            select(Unit).where(Unit.project_id == project_id)
        )
        units = result.scalars().all()
        return {unit.id: unit for unit in units}

    def _strip_html_tags(self, html_content: str) -> str:
        """移除 HTML 标签，获取纯文本"""
        if not html_content:
            return ""
        import re
        # 移除 HTML 标签
        text = re.sub(r'<[^>]+>', ' ', html_content)
        # 解码 HTML 实体
        from html import unescape
        text = unescape(text)
        # 清理多余空白
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    async def export_to_txt(self, project_id: str) -> AsyncIterator[bytes]:
        """导出为 TXT 格式"""
        project = await self.get_project_info(project_id)
        if not project:
            raise ValueError("项目不存在")

        chapters = await self.get_project_chapters(project_id)
        units_map = await self.get_units_map(project_id)

        # 构建 TXT 内容
        lines = []
        lines.append(f"《{project.title}》")
        lines.append("")
        if project.description:
            lines.append(f"简介：{project.description}")
            lines.append("")
        if project.genre:
            lines.append(f"类型：{project.genre}")
            lines.append("")
        lines.append("=" * 50)
        lines.append("")

        current_unit_id = None
        for chapter in chapters:
            # 如果是单元剧且有单元信息，显示单元标题
            if chapter.unit_id and chapter.unit_id != current_unit_id:
                unit = units_map.get(chapter.unit_id)
                if unit:
                    lines.append("")
                    lines.append(f"第{unit.unit_number}单元 {unit.title or ''}")
                    lines.append("-" * 30)
                    current_unit_id = chapter.unit_id

            lines.append("")
            lines.append(f"第{chapter.chapter_number}章 {chapter.title or ''}")
            lines.append("")

            # 移除 HTML 标签，获取纯文本内容
            content = self._strip_html_tags(chapter.content or "")
            lines.append(content)
            lines.append("")
            lines.append(f"（字数：{chapter.word_count or 0}）")
            lines.append("")
            lines.append("=" * 50)
            lines.append("")

        # 统计信息
        total_words = sum(chapter.word_count or 0 for chapter in chapters)
        lines.append("")
        lines.append("=" * 50)
        lines.append(f"全书共{len(chapters)}章，总计{total_words}字")
        lines.append(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # 生成文件
        content = "\n".join(lines)
        yield content.encode('utf-8')

    async def export_to_epub(self, project_id: str) -> AsyncIterator[bytes]:
        """导出为 EPUB 格式"""
        project = await self.get_project_info(project_id)
        if not project:
            raise ValueError("项目不存在")

        chapters = await self.get_project_chapters(project_id)
        units_map = await self.get_units_map(project_id)

        # 创建 EPUB 书籍
        book = epub.EpubBook()

        # 设置元数据
        book.set_identifier(f"inkmate-{project_id}-{datetime.now().strftime('%Y%m%d%H%M%S')}")
        book.set_title(project.title)
        book.set_language('zh-CN')
        book.add_author('InkMate')

        if project.description:
            book.add_metadata('DC', 'description', project.description)

        # 创建 CSS 样式
        style = '''
            body {
                font-family: "Microsoft YaHei", "SimSun", serif;
                line-height: 1.8;
                margin: 10%;
            }
            h1 {
                text-align: center;
                font-size: 1.5em;
                margin: 2em 0 1em 0;
            }
            h2 {
                text-align: center;
                font-size: 1.2em;
                margin: 1.5em 0 0.5em 0;
            }
            p {
                text-indent: 2em;
                margin: 0.5em 0;
            }
            .chapter-title {
                text-align: center;
                font-size: 1.3em;
                margin: 1em 0;
            }
            .unit-title {
                text-align: center;
                font-size: 1.1em;
                color: #666;
                margin: 2em 0 1em 0;
            }
            .word-count {
                text-align: right;
                font-size: 0.8em;
                color: #999;
                margin-top: 1em;
            }
        '''
        css = epub.EpubItem(
            uid="style",
            file_name="style/main.css",
            media_type="text/css",
            content=style
        )
        book.add_item(css)

        # 创建章节
        epub_chapters = []
        toc = []
        spine = ['nav']

        current_unit_id = None
        for chapter in chapters:
            # 如果是单元剧且有单元信息，创建单元分隔页
            if chapter.unit_id and chapter.unit_id != current_unit_id:
                unit = units_map.get(chapter.unit_id)
                if unit:
                    unit_content = f'''
                        <html>
                        <head>
                            <link rel="stylesheet" type="text/css" href="style/main.css"/>
                        </head>
                        <body>
                            <h1 class="unit-title">第{unit.unit_number}单元 {unit.title or ''}</h1>
                        </body>
                        </html>
                    '''
                    unit_chapter = epub.EpubHtml(
                        title=f"第{unit.unit_number}单元",
                        file_name=f"unit_{unit.unit_number}.xhtml",
                        lang='zh-CN'
                    )
                    unit_chapter.content = unit_content
                    unit_chapter.add_item(css)
                    book.add_item(unit_chapter)
                    spine.append(unit_chapter)
                    current_unit_id = chapter.unit_id

            # 创建章节内容
            chapter_content = f'''
                <html>
                <head>
                    <link rel="stylesheet" type="text/css" href="style/main.css"/>
                </head>
                <body>
                    <h1 class="chapter-title">第{chapter.chapter_number}章 {chapter.title or ''}</h1>
                    {chapter.content or ''}
                    <p class="word-count">（字数：{chapter.word_count or 0}）</p>
                </body>
                </html>
            '''

            epub_chapter = epub.EpubHtml(
                title=chapter.title or f"第{chapter.chapter_number}章",
                file_name=f"chapter_{chapter.chapter_number}.xhtml",
                lang='zh-CN'
            )
            epub_chapter.content = chapter_content
            epub_chapter.add_item(css)

            book.add_item(epub_chapter)
            spine.append(epub_chapter)
            toc.append(epub_chapter)
            epub_chapters.append(epub_chapter)

        # 设置目录和书脊
        book.toc = toc
        book.spine = spine

        # 添加导航文件
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # 生成 EPUB 文件到临时缓冲区
        buffer = io.BytesIO()
        epub.write_epub(buffer, book, {})
        buffer.seek(0)
        yield buffer.read()

    async def export_to_docx(self, project_id: str) -> AsyncIterator[bytes]:
        """导出为 DOCX 格式"""
        project = await self.get_project_info(project_id)
        if not project:
            raise ValueError("项目不存在")

        chapters = await self.get_project_chapters(project_id)
        units_map = await self.get_units_map(project_id)

        # 创建 Word 文档
        doc = Document()

        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)

        # 添加标题
        title = doc.add_heading(project.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加简介
        if project.description:
            doc.add_paragraph(f"简介：{project.description}")
        if project.genre:
            doc.add_paragraph(f"类型：{project.genre}")

        doc.add_paragraph("=" * 50)

        current_unit_id = None
        for chapter in chapters:
            # 如果是单元剧且有单元信息，显示单元标题
            if chapter.unit_id and chapter.unit_id != current_unit_id:
                unit = units_map.get(chapter.unit_id)
                if unit:
                    unit_heading = doc.add_heading(
                        f"第{unit.unit_number}单元 {unit.title or ''}",
                        level=2
                    )
                    unit_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_unit_id = chapter.unit_id

            # 添加章节标题
            chapter_title = doc.add_heading(
                f"第{chapter.chapter_number}章 {chapter.title or ''}",
                level=3
            )
            chapter_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 添加章节内容（移除 HTML 标签）
            content = self._strip_html_tags(chapter.content or "")
            if content:
                for paragraph in content.split('\n'):
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph)
                        p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进 2 字符

            # 添加字数统计
            word_count_p = doc.add_paragraph(f"（字数：{chapter.word_count or 0}）")
            word_count_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            word_count_p.style.font.size = Pt(9)
            word_count_p.style.font.color.rgb = RGBColor(150, 150, 150)

            doc.add_paragraph("=" * 50)

        # 添加统计信息
        total_words = sum(chapter.word_count or 0 for chapter in chapters)
        doc.add_paragraph(f"全书共{len(chapters)}章，总计{total_words}字")
        doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # 生成文件到缓冲区
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        yield buffer.read()
